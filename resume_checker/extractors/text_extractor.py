"""Text extraction module for various file formats."""

import logging
from pathlib import Path
from typing import Optional, Dict, Any
import io

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TextExtractor:
    """Extracts text from various file formats including PDF, DOCX, and TXT."""
    
    def __init__(self):
        """Initialize the text extractor."""
        self.logger = logging.getLogger(__name__)
        
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension == '.docx':
                return self._extract_from_docx(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file using multiple methods."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available. Install PyPDF2 and pdfplumber.")
            
        text = ""
        pages = 0
        method_used = "unknown"
        
        # Try pdfplumber first (better text extraction)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                pages = len(pdf.pages)
                method_used = "pdfplumber"
                
        except Exception as e:
            self.logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    pages = len(pdf_reader.pages)
                    method_used = "PyPDF2"
                    
            except Exception as e:
                self.logger.error(f"PyPDF2 also failed: {str(e)}")
                raise
        
        return {
            "text": text.strip(),
            "pages": pages,
            "file_type": "pdf",
            "extraction_method": method_used,
            "file_size": file_path.stat().st_size
        }
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing library not available. Install python-docx.")
            
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            
            return {
                "text": text.strip(),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_type": "docx",
                "extraction_method": "python-docx",
                "file_size": file_path.stat().st_size
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting from DOCX: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        
                    return {
                        "text": text.strip(),
                        "lines": len(text.splitlines()),
                        "file_type": "txt",
                        "extraction_method": f"text_file ({encoding})",
                        "file_size": file_path.stat().st_size,
                        "encoding": encoding
                    }
                    
                except UnicodeDecodeError:
                    continue
                    
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            self.logger.error(f"Error extracting from TXT: {str(e)}")
            raise
    
    def extract_from_string(self, text: str, file_type: str = "text") -> Dict[str, Any]:
        """
        Extract text from a string (useful for testing or direct text input).
        
        Args:
            text: The text content
            file_type: Type identifier for the content
            
        Returns:
            Dictionary containing text and metadata
        """
        return {
            "text": text.strip(),
            "lines": len(text.splitlines()),
            "file_type": file_type,
            "extraction_method": "direct_string",
            "file_size": len(text.encode('utf-8'))
        }