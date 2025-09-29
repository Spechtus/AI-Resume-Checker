import io
from app.services.extraction import extract_text_from_file


def test_extract_txt_utf8():
    content = "Hello résumé checker!".encode("utf-8")
    text = extract_text_from_file(content, "sample.txt")
    assert "résumé" in text


def test_extract_docx(tmp_path):
    try:
        from docx import Document
    except Exception:  # pragma: no cover
        assert False, "python-docx not installed"

    docx_file = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("John Doe")
    document.add_paragraph("Experience: 5 years in Python")
    document.save(docx_file)

    data = docx_file.read_bytes()
    text = extract_text_from_file(data, "resume.docx")
    assert "John Doe" in text
    assert "Python" in text


def test_extract_pdf(tmp_path):
    # Create a simple PDF using reportlab if available; otherwise skip
    try:
        from reportlab.pdfgen import canvas
    except Exception:
        # Fallback: ensure pdfminer dependency exists by asserting import
        import pdfminer  # noqa: F401
        return

    pdf_path = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "Jane Doe")
    c.drawString(100, 730, "Skills: FastAPI, NLP")
    c.save()

    data = pdf_path.read_bytes()
    text = extract_text_from_file(data, "resume.pdf")
    assert "Jane Doe" in text
    assert "FastAPI" in text

